import os
import json
import requests
import docx
import pandas as pd
import numpy as np
from flask import Flask, render_template, jsonify
import urllib.request
import bs4
from bs4 import BeautifulSoup


app = Flask(__name__)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/products")
def products():
    doc = docx.Document("static/YuBambu.docx") ## document with products information
    df = pd.DataFrame()
    tables = doc.tables[0]
# Getting the original data from the document to a list
    ls = []
    product_amount = 0
    for row in tables.rows:
        # simply find the amount of product in the list
        product_amount = product_amount + 1
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                ls.append(paragraph.text)

    ls = list(filter(None, ls))  # this removes the empty strings from the list
    ls = ls[7:]

    product_dict = {  # this is the dict that my data for the site is stored in
        "Product Name": [],
        "Product URL": [],
        "Product Price": [],
        "Manufacturer Price": [],
        "Shipping Price": [],
        "Shipping Times": [],
        "Category": [],
    }

    counter = 0
    for i in range(len(ls)):
        if counter == 0:
            product_dict["Product Name"].append(ls[i])

        if counter == 1:
            product_dict["Product URL"].append(ls[i])

        if counter == 2:
            product_dict["Product Price"].append(ls[i])

        if counter == 3:
            product_dict["Manufacturer Price"].append(ls[i])

        if counter == 4:
            product_dict["Shipping Price"].append(ls[i])

        if counter == 5:
            product_dict["Shipping Times"].append(ls[i])

        if counter == 7:
            product_dict["Category"].append(ls[i])
            counter = -1

        counter = counter + 1

    # section above appends the value of the list to the corresponding key value in the dict
    url_substring = []
    for i in range(product_amount - 1):
        url_string = product_dict["Product URL"][i]
        url_split_string = url_string.split("?", 1)
        url_substring.append(url_split_string[0])
    ## section above splits the url into a useable url i.e remove the junk 

    product_dict["Product URL"] = []
    for i in range(product_amount - 1):
        product_dict["Product URL"].append(url_substring[i])
    # section above cleans the url from the folder by creating a substring list of strings, then the product dict is cleared, then updated

    final_img_src_list = []
    for i in range(product_amount - 1):
        grab_url_html = urllib.request.urlopen(product_dict["Product URL"][i])
        grabbed_url_html = grab_url_html.read()
        # above grabs the url html code from product dict url and stores the variable 
        decode_grabbed_url_html = grabbed_url_html.decode(
            "utf8")  # decodes the data into url
        grab_url_html.close()

        # section above grabs the url html code from the link provided in the product_dict url key
        html = decode_grabbed_url_html
        soup = BeautifulSoup(html, "lxml")
        img_url_first_script = soup.find('script')
        # section above finds the first element with the script tag, this was selected due to manual selection
        soup_string = str(img_url_first_script)
        split = soup_string.split("https://", 1)
        converted = "https://" + split[1]
        converted_split = converted.split(".jpg", 1)
        final_img_src = converted_split[0] + ".jpg"
        final_img_src_list.append(final_img_src)
        # this converts the BS4 element into a string, then chops the string into the useful images, then appends to list

    review_text_string = []
    for i in range(product_amount -1):
        headers = {"apikey": "59681810-5262-11eb-8245-47538fc8149a"}
        params = (
        ("url", product_dict["Product URL"][i]), ("amount", "1"))
        response = requests.get('https://app.reviewapi.io/api/v1/reviews', headers=headers, params=params)
        review_data_json = response.json()
        review_data = review_data_json
        #review_text_string.append(review_data['reviews'][0]['text'])
        review_text_string = review_data # the code above is commented out as api calls reached limit so this line is added to stop crashing
    ##section above grabs the product url from product dict and passed that data through the review api
    
    display_cards_rows = 0
    display_cards_rows = (product_amount / 4)
    x = 0
    z = 4

    return render_template("products.html", x=x, z=z, display_cards_rows=int(display_cards_rows), product_dict=product_dict,  product_amount=product_amount, page=final_img_src_list, pagetest=decode_grabbed_url_html, review_text_string=review_text_string)

@app.route("/test")
def test():
    
    return render_template("test.html")


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/contact")
def contact():
    return render_template("contact.html")


if __name__ == "__main__":
    app.run(
        host=os.environ.get("IP", "0.0.0.0"),
        port=int(os.environ.get("PORT", "5000")),
        debug=True)
